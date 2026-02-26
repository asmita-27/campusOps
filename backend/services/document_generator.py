"""
Document Generator Service
Creates downloadable documents (DOCX) from generated content
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import re


class DocumentGenerator:
    """Service for generating formatted DOCX documents"""
    
    def __init__(self):
        """Initialize Document Generator"""
        # Get the backend directory path
        backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.output_folder = os.path.join(backend_dir, 'outputs', 'documents')
        os.makedirs(self.output_folder, exist_ok=True)
    
    def generate_event_document(self, content, document_type='event_plan', metadata=None):
        """
        Generate a formatted DOCX document from event content
        
        Args:
            content (str or dict): The generated event content (string or dict with 'content' key)
            document_type (str): Type of document (event_plan, summary, report)
            metadata (dict): Additional metadata (event_description, etc.)
        
        Returns:
            dict: Result with file path and success status
        """
        try:
            # Handle dict input (from LLM service)
            if isinstance(content, dict):
                content = content.get('content', str(content))
            
            # Ensure content is a string
            if not isinstance(content, str):
                content = str(content)
            
            # Create a new Document
            doc = Document()
            
            # Set default styles
            self._set_document_styles(doc)
            
            # Check if content follows the form-style template (table-based)
            if '[TABLE:' in content or 'Name of the Club' in content:
                # Form-style report - parse as structured tables
                self._parse_form_style_content(doc, content, metadata.get('image_paths'))
            else:
                # Standard markdown report - use old parsing
                # Add title
                title = self._get_title_from_type(document_type)
                
                # Add main heading without "EVENT"
                title_para = doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add "Event Report/Plan/Summary" subheading
                subtitle = self._get_subtitle_from_type(document_type)
                subtitle_para = doc.add_heading(subtitle, level=1)
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add metadata section
                if metadata:
                    self._add_metadata_section(doc, metadata)
                
                # Parse and add content (skip separator, it's in content)
                self._parse_and_add_content(doc, content, metadata.get('image_paths'))
            
            # Add footer with generation timestamp
            self._add_footer(doc)
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_type = document_type.replace(' ', '_').lower()
            filename = f"{safe_type}_{timestamp}.docx"
            
            # Ensure output folder exists before saving
            os.makedirs(self.output_folder, exist_ok=True)
            filepath = os.path.join(self.output_folder, filename)
            
            # Save document
            doc.save(filepath)
            
            print(f"Document saved successfully to: {filepath}")
            
            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'message': 'Document generated successfully'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to generate document'
            }
    
    def _set_document_styles(self, doc):
        """Set default styles for the document"""
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _get_title_from_type(self, document_type):
        """Get formatted title based on document type"""
        # This will be extracted from the content (e.g., "Event Report of Club/Committee FF 984")
        # For now, return a generic title
        return f'Event {document_type.replace("_", " ").title()}'
    
    def _get_subtitle_from_type(self, document_type):
        """Get subtitle based on document type"""
        subtitles = {
            'event_plan': 'Event Plan',
            'summary': 'Event Summary',
            'report': 'Event Report'
        }
        return subtitles.get(document_type, 'Event Document')
    
    def _add_metadata_section(self, doc, metadata):
        """Add metadata information to document"""
        doc.add_paragraph()
        
        # Add metadata table or formatted info
        if metadata.get('event_description'):
            p = doc.add_paragraph()
            p.add_run('Event Description: ').bold = True
            p.add_run(metadata['event_description'][:200] + '...' if len(metadata['event_description']) > 200 else metadata['event_description'])
        
        # Add generation date
        p = doc.add_paragraph()
        p.add_run('Generated: ').bold = True
        p.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))
        
        if metadata.get('template_used'):
            p = doc.add_paragraph()
            p.add_run('Template: ').bold = True
            p.add_run(metadata.get('template_format', 'Custom Template'))
        
        doc.add_paragraph()
    
    def _parse_and_add_content(self, doc, content, image_paths=None):
        """
        Parse the AI-generated content and format it properly
        Handles headers, bullet points, tables, sections, images, etc.
        
        Args:
            doc: Document object
            content: Generated content text
            image_paths: List of image file paths to insert
        """
        # Split content into lines
        lines = content.split('\n')
        
        # Track which images have been added
        geo_images_added = False
        non_geo_images_added = False
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add spacing
                doc.add_paragraph()
                i += 1
                continue
            
            # Check if this is the start of a markdown table
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                # Parse the table
                table_lines = [line]
                i += 1
                # Skip the separator line
                i += 1
                # Collect remaining table rows
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Add the table to document
                self._add_markdown_table(doc, table_lines)
                continue
            
            # Check if line is a header (all caps or starts with numbers like "1.", "2.")
            if self._is_header(line):
                heading_level = self._get_heading_level(line)
                doc.add_heading(line, level=heading_level)
                
                # Check if this is a photo section header
                if image_paths and 'photograph section' in line.lower():
                    if 'geo-tagged' in line.lower() and not geo_images_added:
                        # Add first half of images for geo-tagged section
                        self._add_images_to_document(doc, image_paths[:len(image_paths)//2] if len(image_paths) > 1 else image_paths)
                        geo_images_added = True
                    elif 'non geo-tagged' in line.lower() and not non_geo_images_added:
                        # Add second half of images for non-geo-tagged section
                        self._add_images_to_document(doc, image_paths[len(image_paths)//2:] if len(image_paths) > 1 else [])
                        non_geo_images_added = True
            
            # Check if line is a bullet point
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                bullet_text = line[1:].strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
            
            # Check if line is a numbered item
            elif self._is_numbered_item(line):
                doc.add_paragraph(line, style='List Number')
            
            # Check if line contains bold markers (e.g., **text**)
            elif '**' in line:
                self._add_formatted_paragraph(doc, line)
            
            # Regular paragraph
            else:
                doc.add_paragraph(line)
            
            i += 1
    
    def _parse_form_style_content(self, doc, content, image_paths=None):
        """
        Parse form-style content with table-based structure
        
        Args:
            doc: Document object
            content: Generated content text with [TABLE:...] markers
            image_paths: List of image file paths to insert
        """
        lines = content.split('\n')
        i = 0
        geo_images_added = False
        non_geo_images_added = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Handle title line (first line with "Title:")
            if line.startswith('Title:'):
                title_text = line.replace('Title:', '').strip()
                title_para = doc.add_paragraph(title_text)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.runs[0].bold = True
                title_para.runs[0].font.size = Pt(14)
                doc.add_paragraph()  # Add spacing
                i += 1
                continue
            
            # Handle table markers
            if line.startswith('[TABLE:'):
                # Extract table description
                table_desc = line.replace('[TABLE:', '').replace(']', '').strip()
                i += 1
                
                # Collect table rows until we hit next section or table marker
                table_rows = []
                while i < len(lines):
                    row_line = lines[i].strip()
                    
                    # Stop if we hit another table marker, empty line followed by marker, or section header
                    if (row_line.startswith('[TABLE:') or 
                        row_line.startswith('##') or
                        (not row_line and i + 1 < len(lines) and 
                         (lines[i+1].strip().startswith('[TABLE:') or lines[i+1].strip().startswith('##')))):
                        break
                    
                    if row_line and '|' in row_line:
                        table_rows.append(row_line)
                    
                    i += 1
                
                # Create the table
                if table_rows:
                    if 'Program Outcomes' in table_desc:
                        # Special handling for Program Outcomes table (4 columns)
                        self._add_program_outcomes_table(doc, table_rows)
                    else:
                        # Standard 2-column field/value table
                        self._add_field_value_table(doc, table_rows)
                    
                    doc.add_paragraph()  # Add spacing after table
                continue
            
            # Handle section headers (##, ###)
            if line.startswith('##'):
                header_text = line.replace('##', '').replace('###', '').strip()
                doc.add_heading(header_text, level=2)
                
                # Check if this is a photo section header
                if image_paths and 'photograph' in header_text.lower():
                    if 'geo-tagged' in header_text.lower() and not geo_images_added:
                        # Add first half of images
                        self._add_images_to_document(doc, image_paths[:len(image_paths)//2] if len(image_paths) > 1 else image_paths)
                        geo_images_added = True
                    elif 'non geo-tagged' in header_text.lower() and not non_geo_images_added:
                        # Add second half of images
                        self._add_images_to_document(doc, image_paths[len(image_paths)//2:] if len(image_paths) > 1 else [])
                        non_geo_images_added = True
                i += 1
                continue
            
            # Handle bullet points
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                bullet_text = line[1:].strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
                i += 1
                continue
            
            # Regular paragraph
            doc.add_paragraph(line)
            i += 1
    
    def _add_field_value_table(self, doc, table_rows):
        """
        Add a 2-column field/value table (for event details)
        
        Args:
            doc: Document object
            table_rows: List of "Field | Value" strings
        """
        if not table_rows:
            return
        
        # Create table with 2 columns
        table = doc.add_table(rows=len(table_rows), cols=2)
        table.style = 'Table Grid'
        
        # Populate table
        for idx, row_text in enumerate(table_rows):
            parts = [p.strip() for p in row_text.split('|') if p.strip()]
            if len(parts) >= 2:
                # Field name (left column)
                cell_0 = table.rows[idx].cells[0]
                cell_0.text = parts[0]
                cell_0.paragraphs[0].runs[0].bold = True
                
                # Apply light green shading to certain rows (optional - customize as needed)
                if any(keyword in parts[0].lower() for keyword in ['club', 'event', 'student', 'mode', 'participants']):
                    self._shade_cell(cell_0, 'D6EAD6')  # Light green
                
                # Value (right column)
                cell_1 = table.rows[idx].cells[1]
                cell_1.text = parts[1]
    
    def _add_program_outcomes_table(self, doc, table_rows):
        """
        Add a 4-column Program Outcomes table with ratings
        
        Args:
            doc: Document object
            table_rows: List of table rows with "S.No. | Program Outcome | Rating (0-3) | Remarks"
        """
        if not table_rows:
            return
        
        # First row is header
        header_parts = [p.strip() for p in table_rows[0].split('|') if p.strip()]
        num_cols = len(header_parts)
        
        # Create table
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        
        # Add header row with formatting
        for col_idx, header_text in enumerate(header_parts):
            cell = table.rows[0].cells[col_idx]
            cell.text = header_text
            cell.paragraphs[0].runs[0].bold = True
            self._shade_cell(cell, 'D6EAD6')  # Light green header
        
        # Add data rows
        for row_idx in range(1, len(table_rows)):
            row_parts = [p.strip() for p in table_rows[row_idx].split('|') if p.strip()]
            for col_idx in range(min(len(row_parts), num_cols)):
                table.rows[row_idx].cells[col_idx].text = row_parts[col_idx]
    
    def _shade_cell(self, cell, color_hex):
        """
        Apply background color shading to a table cell
        
        Args:
            cell: Table cell object
            color_hex: Hex color code (e.g., 'D6EAD6' for light green)
        """
        try:
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            # If shading fails, just log and continue
            print(f"Warning: Could not apply cell shading: {e}")
    
    def _add_markdown_table(self, doc, table_lines):
        """
        Add a markdown table to the document with professional formatting
        
        Args:
            doc: The document object
            table_lines: List of table lines in markdown format
        """
        if not table_lines:
            return
        
        try:
            # Parse header row
            header_row = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
            num_cols = len(header_row)
            
            # Parse data rows
            data_rows = []
            for line in table_lines[1:]:
                row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                if row_data:
                    data_rows.append(row_data)
            
            if not data_rows:
                return
            
            # Create table
            num_rows = len(data_rows) + 1  # +1 for header
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            # Add header row with bold and background
            header_cells = table.rows[0].cells
            for i, header_text in enumerate(header_row):
                if i < num_cols:
                    # Remove markdown bold markers
                    header_text = header_text.replace('**', '')
                    cell = header_cells[i]
                    cell.text = header_text
                    
                    # Make header bold with larger font
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                    
                    # Add shading to header (light blue-gray)
                    from docx.oxml.shared import OxmlElement
                    from docx.oxml.ns import qn
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D9E2F3')  # Light blue background
                    cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Add data rows
            for row_idx, row_data in enumerate(data_rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        # Remove markdown bold markers
                        cell_text = cell_text.replace('**', '')
                        cell = row_cells[col_idx]
                        cell.text = cell_text
                        
                        # Set font size for data cells
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # Add spacing after table
            doc.add_paragraph()
            
        except Exception as e:
            print(f"Error adding table: {e}")
            # Fallback: add as paragraph
            for line in table_lines:
                doc.add_paragraph(line)
    
    def _is_header(self, line):
        """Check if line should be treated as a header"""
        # All caps with at least 3 words
        if line.isupper() and len(line.split()) >= 2:
            return True
        
        # Starts with number pattern like "1.", "1.1", etc.
        if re.match(r'^\d+\.(\d+\.?)?\s+[A-Z]', line):
            return True
        
        # Ends with colon and is relatively short
        if line.endswith(':') and len(line) < 60:
            return True
        
        return False
    
    def _get_heading_level(self, line):
        """Determine appropriate heading level"""
        if line.isupper() and len(line.split()) <= 5:
            return 1
        if re.match(r'^\d+\.\s+', line):
            return 2
        if re.match(r'^\d+\.\d+\.\s+', line):
            return 3
        return 2
    
    def _is_numbered_item(self, line):
        """Check if line is a numbered list item"""
        return re.match(r'^\d+[\.\)]\s+', line) is not None
    
    def _add_formatted_paragraph(self, doc, line):
        """Add paragraph with inline formatting (bold, italic, etc.)"""
        p = doc.add_paragraph()
        
        # Split by bold markers
        parts = re.split(r'\*\*(.*?)\*\*', line)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                p.add_run(part)
            else:
                # Bold text
                p.add_run(part).bold = True
    
    def _add_images_to_document(self, doc, image_paths):
        """
        Add images to the document
        
        Args:
            doc: Document object
            image_paths: List of image file paths to insert
        """
        if not image_paths:
            return
        
        try:
            for img_path in image_paths:
                if os.path.exists(img_path):
                    # Add paragraph for spacing
                    doc.add_paragraph()
                    
                    # Add the image with a reasonable width (6 inches)
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    try:
                        # Add image with max width of 6 inches
                        run.add_picture(img_path, width=Inches(6))
                        
                        # Center align the image
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption
                        caption = doc.add_paragraph()
                        caption_run = caption.add_run(f"Image: {os.path.basename(img_path)}")
                        caption_run.font.size = Pt(9)
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing after image
                        doc.add_paragraph()
                        
                    except Exception as e:
                        print(f"Error adding image {img_path}: {e}")
                        # Add note that image couldn't be loaded
                        p = doc.add_paragraph()
                        p.add_run(f"[Image: {os.path.basename(img_path)} - Could not be loaded]").italic = True
                
                else:
                    print(f"Image not found: {img_path}")
                    
        except Exception as e:
            print(f"Error adding images to document: {e}")
    
    def _add_footer(self, doc):
        """Add footer with generation info"""
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f'\nGenerated by CampusOps Event Report Generator\n'
            f'Date: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def cleanup_old_files(self, days=7):
        """
        Clean up old generated files
        
        Args:
            days (int): Delete files older than this many days
        """
        try:
            import time
            now = time.time()
            cutoff = now - (days * 86400)
            
            for filename in os.listdir(self.output_folder):
                filepath = os.path.join(self.output_folder, filename)
                if os.path.isfile(filepath):
                    file_modified = os.path.getmtime(filepath)
                    if file_modified < cutoff:
                        os.remove(filepath)
                        print(f"Deleted old file: {filename}")
        
        except Exception as e:
            print(f"Error cleaning up old files: {e}")
