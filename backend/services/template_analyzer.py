"""
Template Analyzer Service
Extracts structure and formatting from template documents
"""
import os
from docx import Document
from PyPDF2 import PdfReader
import re


class TemplateAnalyzer:
    """Analyzes document templates to extract structure and format"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.docx', '.pdf']
    
    def analyze_template(self, file_path):
        """
        Analyze a template file and extract its structure
        
        Args:
            file_path: Path to the template file
            
        Returns:
            dict: Template analysis with structure and formatting details
        """
        if not os.path.exists(file_path):
            return {
                'error': 'Template file not found',
                'structure': None
            }
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            return self._analyze_text_template(file_path)
        elif file_ext == '.docx':
            return self._analyze_docx_template(file_path)
        elif file_ext == '.pdf':
            return self._analyze_pdf_template(file_path)
        else:
            return {
                'error': f'Unsupported file format: {file_ext}',
                'structure': None
            }
    
    def _analyze_text_template(self, file_path):
        """Analyze plain text template"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            structure = self._extract_structure(content)
            
            return {
                'success': True,
                'format': 'text',
                'content': content,
                'structure': structure,
                'metadata': {
                    'line_count': len(content.split('\n')),
                    'word_count': len(content.split()),
                    'has_sections': len(structure['sections']) > 0
                }
            }
        except Exception as e:
            return {
                'error': f'Error reading text file: {str(e)}',
                'structure': None
            }
    
    def _analyze_docx_template(self, file_path):
        """Analyze DOCX template"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            full_text = []
            headings = []
            paragraphs_data = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
                    
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        headings.append({
                            'text': text,
                            'level': para.style.name,
                            'is_bold': para.runs[0].bold if para.runs else False
                        })
                    
                    paragraphs_data.append({
                        'text': text,
                        'style': para.style.name,
                        'is_bold': para.runs[0].bold if para.runs else False,
                        'alignment': str(para.alignment) if para.alignment else 'LEFT'
                    })
            
            # Extract tables structure
            tables_structure = []
            for table in doc.tables:
                table_data = {
                    'rows': len(table.rows),
                    'cols': len(table.columns),
                    'headers': [cell.text for cell in table.rows[0].cells]
                }
                tables_structure.append(table_data)
            
            combined_text = '\n'.join(full_text)
            structure = self._extract_structure(combined_text)
            
            return {
                'success': True,
                'format': 'docx',
                'content': combined_text,
                'structure': structure,
                'formatting': {
                    'headings': headings,
                    'paragraphs': paragraphs_data[:5],  # First 5 paragraphs for reference
                    'tables': tables_structure
                },
                'metadata': {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'heading_count': len(headings)
                }
            }
        except Exception as e:
            return {
                'error': f'Error reading DOCX file: {str(e)}',
                'structure': None
            }
    
    def _analyze_pdf_template(self, file_path):
        """Analyze PDF template"""
        try:
            reader = PdfReader(file_path)
            
            # Extract text from all pages
            full_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
            
            combined_text = '\n'.join(full_text)
            structure = self._extract_structure(combined_text)
            
            return {
                'success': True,
                'format': 'pdf',
                'content': combined_text,
                'structure': structure,
                'metadata': {
                    'page_count': len(reader.pages),
                    'word_count': len(combined_text.split())
                }
            }
        except Exception as e:
            return {
                'error': f'Error reading PDF file: {str(e)}',
                'structure': None
            }
    
    def _extract_structure(self, text):
        """
        Extract document structure from text
        Identifies sections, numbering patterns, headings, etc.
        """
        lines = text.split('\n')
        
        # Identify sections (lines that look like headings)
        sections = []
        section_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'^[A-Z\s]{3,}:?$',  # ALL CAPS headings
            r'^\d+\.\s+[A-Z]',   # Numbered sections like "1. Introduction"
            r'^[IVX]+\.\s+[A-Z]',  # Roman numeral sections
            r'^[A-Z][a-z]+\s*:$',  # Title case with colon
        ]
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            for pattern in section_patterns:
                if re.match(pattern, line):
                    sections.append({
                        'line_number': i,
                        'text': line,
                        'pattern': pattern
                    })
                    break
        
        # Identify numbering style
        numbering_style = 'none'
        if any(re.match(r'^\d+\.', line.strip()) for line in lines):
            numbering_style = 'numeric'
        elif any(re.match(r'^[a-z]\)', line.strip()) for line in lines):
            numbering_style = 'alphabetic'
        elif any(re.match(r'^[IVX]+\.', line.strip()) for line in lines):
            numbering_style = 'roman'
        elif any(re.match(r'^\*\s+', line.strip()) or re.match(r'^-\s+', line.strip()) for line in lines):
            numbering_style = 'bullets'
        
        # Identify common section names
        common_sections = []
        section_keywords = [
            'abstract', 'introduction', 'background', 'objective', 'objectives',
            'methodology', 'method', 'approach', 'timeline', 'schedule',
            'budget', 'cost', 'resources', 'team', 'participants',
            'results', 'outcome', 'outcomes', 'conclusion', 'summary',
            'recommendations', 'next steps', 'follow-up', 'references'
        ]
        
        for line in lines:
            line_lower = line.strip().lower()
            for keyword in section_keywords:
                if keyword in line_lower and len(line.strip()) < 50:
                    common_sections.append(line.strip())
                    break
        
        # Detect formatting patterns
        has_date_format = bool(re.search(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}', text))
        has_bullet_points = bool(re.search(r'[•\-\*]\s+\w+', text))
        has_numbered_list = bool(re.search(r'\d+\.\s+\w+', text))
        
        return {
            'sections': sections,
            'common_sections': common_sections,
            'numbering_style': numbering_style,
            'formatting_features': {
                'has_dates': has_date_format,
                'has_bullets': has_bullet_points,
                'has_numbered_lists': has_numbered_list
            },
            'total_sections': len(sections)
        }
    
    def create_template_prompt(self, template_analysis):
        """
        Create a prompt instruction for LLM based on template analysis
        
        Args:
            template_analysis: Output from analyze_template()
            
        Returns:
            str: Detailed formatting instructions for the LLM
        """
        if not template_analysis.get('success'):
            return ""
        
        structure = template_analysis.get('structure', {})
        metadata = template_analysis.get('metadata', {})
        formatting = template_analysis.get('formatting', {})
        
        # Build template instructions
        instructions = []
        
        instructions.append("IMPORTANT: Follow this exact document template structure:")
        instructions.append("")
        
        # Document format
        instructions.append(f"Document Format: {template_analysis.get('format', 'text').upper()}")
        
        # Sections
        if structure.get('common_sections'):
            instructions.append("\nRequired Sections:")
            for section in structure['common_sections'][:10]:  # Limit to 10 sections
                instructions.append(f"  - {section}")
        
        # Numbering style
        numbering = structure.get('numbering_style', 'none')
        if numbering != 'none':
            instructions.append(f"\nNumbering Style: {numbering.upper()}")
            if numbering == 'numeric':
                instructions.append("  Use: 1. 2. 3. for main points")
            elif numbering == 'alphabetic':
                instructions.append("  Use: a) b) c) for sub-points")
            elif numbering == 'roman':
                instructions.append("  Use: I. II. III. for sections")
            elif numbering == 'bullets':
                instructions.append("  Use: • or - for bullet points")
        
        # Formatting features
        features = structure.get('formatting_features', {})
        if features.get('has_dates'):
            instructions.append("\nInclude dates in format: DD/MM/YYYY or similar")
        if features.get('has_bullets'):
            instructions.append("Use bullet points for lists")
        if features.get('has_numbered_lists'):
            instructions.append("Use numbered lists for sequential items")
        
        # DOCX specific
        if template_analysis.get('format') == 'docx' and formatting:
            headings = formatting.get('headings', [])
            if headings:
                instructions.append("\nHeading Styles:")
                for heading in headings[:5]:  # First 5 headings
                    instructions.append(f"  {heading['level']}: {heading['text'][:50]}...")
            
            tables = formatting.get('tables', [])
            if tables:
                instructions.append(f"\nInclude {len(tables)} table(s) with appropriate structure")
        
        # Sample content reference
        if template_analysis.get('content'):
            sample = template_analysis['content'][:500]  # First 500 chars
            instructions.append(f"\nTemplate Style Reference:")
            instructions.append(f"```\n{sample}\n```")
        
        return '\n'.join(instructions)
