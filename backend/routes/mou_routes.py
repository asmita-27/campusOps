"""
MOU (Memorandum of Understanding) Generator Routes
AI-powered MOU document generation
"""

from flask import Blueprint, request, jsonify, current_app, send_file
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

bp = Blueprint('mou', __name__, url_prefix='/api/mou')


@bp.route('/generate', methods=['POST'])
def generate_mou():
    """
    Generate MOU document using AI
    
    Expects:
        - party1_name: First party name (usually the club)
        - party1_address: First party address
        - party2_name: Second party name (sponsor/partner)
        - party2_address: Second party address
        - purpose: Purpose of the agreement
        - event_name: Name of the event (optional)
        - duration: Duration of agreement
        - terms: Additional terms (optional)
    """
    try:
        data = request.json
        
        # Validate required fields
        required_fields = ['party1_name', 'party2_name', 'purpose']
        for field in required_fields:
            if not data.get(field):
                return jsonify({
                    'success': False,
                    'error': f'Missing required field: {field}'
                }), 400
        
        party1_name = data['party1_name']
        party1_address = data.get('party1_address', '')
        party2_name = data['party2_name']
        party2_address = data.get('party2_address', '')
        purpose = data['purpose']
        event_name = data.get('event_name', '')
        duration = data.get('duration', '1 year')
        additional_terms = data.get('terms', '')
        
        # Build AI prompt
        prompt = f"""Generate a professional Memorandum of Understanding (MOU) between:

PARTY 1 (First Party):
Name: {party1_name}
Address: {party1_address}

PARTY 2 (Second Party):
Name: {party2_name}
Address: {party2_address}

PURPOSE: {purpose}
{f"EVENT: {event_name}" if event_name else ""}
DURATION: {duration}
{f"ADDITIONAL TERMS: {additional_terms}" if additional_terms else ""}

Generate a complete, professional MOU with the following sections:
1. Preamble (identifying both parties)
2. Purpose and Objectives
3. Scope of Collaboration
4. Roles and Responsibilities
   - Party 1 Obligations
   - Party 2 Obligations
5. Duration and Termination
6. Financial Terms (if applicable)
7. Intellectual Property Rights
8. Confidentiality
9. Dispute Resolution
10. Miscellaneous Provisions

Make it formal, legally sound, and comprehensive. Use proper legal language but keep it clear and understandable."""

        # Call LLM service
        llm = current_app.llm
        if not llm.is_available():
            return jsonify({
                'success': False,
                'error': 'LLM service is not available'
            }), 503
        
        mou_content = llm.generate_response(prompt)
        
        # Save to database
        db = current_app.db.db
        mou_record = {
            'party1_name': party1_name,
            'party1_address': party1_address,
            'party2_name': party2_name,
            'party2_address': party2_address,
            'purpose': purpose,
            'event_name': event_name,
            'duration': duration,
            'content': mou_content,
            'created_at': datetime.utcnow().isoformat(),
            'status': 'draft'
        }
        
        result = db.mou_documents.insert_one(mou_record)
        mou_id = str(result.inserted_id)
        
        # Generate downloadable DOCX
        filename = f"MOU_{party1_name.replace(' ', '_')}_{party2_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = create_mou_document(mou_content, party1_name, party2_name, filename)
        
        return jsonify({
            'success': True,
            'id': mou_id,
            'content': mou_content,
            'download_path': f'/api/mou/download/{mou_id}',
            'filename': filename
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def create_mou_document(content, party1, party2, filename):
    """Create a formatted DOCX document for the MOU"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('MEMORANDUM OF UNDERSTANDING')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f'Between\n{party1}\nand\n{party2}')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Content - split by paragraphs
    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_format = para.paragraph_format
            para_format.line_spacing = 1.15
            para_format.space_after = Pt(6)
    
    # Signature Section
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_para = doc.add_paragraph()
    sig_para.add_run('_' * 50).font.size = Pt(10)
    doc.add_paragraph(f'Authorized Signatory - {party1}')
    doc.add_paragraph(f'Date: _____________________')
    
    doc.add_paragraph()
    
    sig_para2 = doc.add_paragraph()
    sig_para2.add_run('_' * 50).font.size = Pt(10)
    doc.add_paragraph(f'Authorized Signatory - {party2}')
    doc.add_paragraph(f'Date: _____________________')
    
    # Save document
    backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    outputs_dir = os.path.join(backend_dir, 'outputs', 'mou')
    os.makedirs(outputs_dir, exist_ok=True)
    
    filepath = os.path.join(outputs_dir, filename)
    doc.save(filepath)
    
    return filepath


@bp.route('/download/<mou_id>', methods=['GET'])
def download_mou(mou_id):
    """Download MOU document"""
    try:
        from bson import ObjectId
        db = current_app.db.db
        
        mou = db.mou_documents.find_one({'_id': ObjectId(mou_id)})
        if not mou:
            return jsonify({'success': False, 'error': 'MOU not found'}), 404
        
        # Recreate document
        filename = f"MOU_{mou['party1_name'].replace(' ', '_')}_{mou['party2_name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = create_mou_document(
            mou['content'], 
            mou['party1_name'], 
            mou['party2_name'], 
            filename
        )
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@bp.route('/history', methods=['GET'])
def get_mou_history():
    """Get all MOU documents"""
    try:
        db = current_app.db.db
        mous = list(db.mou_documents.find().sort('created_at', -1).limit(20))
        
        # Serialize MongoDB documents
        for mou in mous:
            mou['_id'] = str(mou['_id'])
            # Don't send full content in list, just metadata
            mou.pop('content', None)
        
        return jsonify({
            'success': True,
            'mous': mous
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@bp.route('/<mou_id>', methods=['GET'])
def get_mou(mou_id):
    """Get specific MOU document"""
    try:
        from bson import ObjectId
        db = current_app.db.db
        
        mou = db.mou_documents.find_one({'_id': ObjectId(mou_id)})
        if not mou:
            return jsonify({'success': False, 'error': 'MOU not found'}), 404
        
        mou['_id'] = str(mou['_id'])
        
        return jsonify({
            'success': True,
            'mou': mou
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
